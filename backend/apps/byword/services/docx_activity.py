import math
import os

from django.conf import settings

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from apps.byword.models import (
    Dictionary,
    ScrambleWord,
    WordSearch,
    Music,
)


def remove_vertical_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            # remove esquerda
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "nil")
            tcBorders.append(left)
            # remove direita
            right = OxmlElement("w:right")
            right.set(qn("w:val"), "nil")
            tcBorders.append(right)

def underline_by_translation(word, translation):
    if translation:
        size = (len(translation) * 3) + 2
    else:
        size = (len(word) * 3) + 2

    return "_" * size


def add_subtitle(doc, text):
    p = doc.add_paragraph()

    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def generate_activity_docx(activity):
    lesson = activity.lesson

    doc = Document()

    # =========================
    # TITLE
    # =========================
    title = doc.add_heading(
        f"Lesson {lesson.number} - {lesson.name}",
        level=1
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================
    # ITEMS
    # =========================
    for item in activity.items.order_by("order"):
        model = item.content_type.model_class()

        # =====================================================
        # DICTIONARY
        # =====================================================
        if model == Dictionary:
            words = (
                Dictionary.objects
                .filter(occurrences__lesson=lesson)
                .distinct()
                .order_by("verb_en")
            )
            if not words.exists():
                continue
            add_subtitle(doc, "Vocabulary")
            words = list(words)
            rows = math.ceil(len(words) / 2)
            table = doc.add_table(rows=rows, cols=2)
            table.style = "Table Grid"
            table.autofit = True
            remove_vertical_borders(table)

            for index, word in enumerate(words):
                col = 0 if index < rows else 1
                row = index if index < rows else index - rows
                underline = underline_by_translation(
                    word.verb_en,
                    word.translation
                )
                table.cell(row, col).text = (
                    f"{word.verb_en} {underline}"
                )

        # =====================================================
        # SCRAMBLE
        # =====================================================
        elif model == ScrambleWord:
            scrambles = (
                ScrambleWord.objects
                .filter(lesson=lesson)
            )
            if not scrambles.exists():
                continue
            add_subtitle(doc, "Scramble Words")
            words = []
            for scramble in scrambles:
                words.extend(
                    scramble.texto_embaralhado.split()
                )
            rows = math.ceil(len(words) / 2)
            table = doc.add_table(rows=rows, cols=2)
            table.style = "Table Grid"
            table.autofit = True
            remove_vertical_borders(table)

            for index, word in enumerate(words):
                col = 0 if index < rows else 1
                row = index if index < rows else index - rows
                underline = "_" * ((len(word) * 3) + 2)
                table.cell(row, col).text = (
                    f"{word} {underline}"
                )

        # =====================================================
        # WORDSEARCH
        # =====================================================
        elif model == WordSearch:

            wordsearch = (
                WordSearch.objects
                .filter(lesson=lesson)
                .first()
            )
            if not wordsearch:
                continue
            add_subtitle(doc, "Word Search")
            png_path = os.path.join(
                settings.MEDIA_ROOT,
                "png",
                f"{wordsearch.id}.png"
            )
            if os.path.exists(png_path):
                doc.add_picture(
                    png_path,
                    width=Inches(5.8)
                )

        # =====================================================
        # MUSIC
        # =====================================================
        elif model == Music:
            musics = Music.objects.filter(
                lesson=lesson
            )
            if not musics.exists():
                continue
            add_subtitle(doc, "Music")
            for music in musics:
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = title.add_run(
                    f"{music.title} - {music.author or ''}"
                )
                run.bold = True
                lyrics = doc.add_paragraph(
                    music.lyrics_spaces or ""
                )
                lyrics.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return doc