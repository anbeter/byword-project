from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from apps.byword.services.docx_engine.registry import DOCX_RENDERERS
from apps.byword.models import Dictionary

from apps.byword.services.docx_activity import (
    add_subtitle,
)
from apps.byword.services.docx_engine.transforms import (
    DOCX_TRANSFORMS,
)

def add_subtitle(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)

def render_text(doc, value, config):
    if not value:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(value))
    run.font.size = Pt(14)

def generate_activity_docx_v2(activity):
    print(DOCX_RENDERERS)
    doc = Document()

    lesson = activity.lesson

    title = doc.add_heading(
        f"Lesson {lesson.number} - {lesson.name}",
        level=1
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if title.runs:
        title.runs[0].font.size = Pt(18)

    for item in activity.items.order_by("order"):

        model = item.content_type.model_class()

        # -------------------------
        # QUERYSET
        # -------------------------

        if model == Dictionary:

            queryset = (
                model.objects
                .filter(
                    occurrences__lesson=lesson
                )
                .distinct()
            )

        else:

            if item.object_id:

                queryset = model.objects.filter(
                    id=item.object_id,
                    lesson=lesson
                )

            else:

                queryset = model.objects.filter(
                    lesson=lesson
                )

        if not queryset.exists():
            continue

        # -------------------------
        # SUBTITLE PADRÃO DO MODEL
        # -------------------------

        model_subtitle = getattr(
            model,
            "docx_subtitle",
            None
        )

        # -------------------------
        # DICTIONARY
        # -------------------------

        if model == Dictionary:

            if model_subtitle:

                add_subtitle(
                    doc,
                    model_subtitle
                )

            renderer = DOCX_RENDERERS[
                "dictionary_table"
            ]

            renderer(
                doc,
                queryset,
                {}
            )

            continue

        # -------------------------
        # DEMAIS MODELS
        # -------------------------

        for obj in queryset:

            subtitle = getattr(
                obj,
                "subtitle",
                model_subtitle
            )

            if subtitle:

                add_subtitle(
                    doc,
                    subtitle
                )

            for field_config in model.docx_fields:

                field_name = field_config["field"]

                renderer_name = field_config["renderer"]

                # -------------------------
                # VALUE
                # -------------------------

                if field_name == "__object__":

                    value = obj

                else:

                    value = getattr(
                        obj,
                        field_name,
                        None
                    )

                # -------------------------
                # TRANSFORM
                # -------------------------

                transform_name = field_config.get(
                    "transform"
                )

                if transform_name:

                    transform_function = DOCX_TRANSFORMS[
                        transform_name
                    ]

                    value = transform_function(
                        value
                    )

                # -------------------------
                # RENDERER
                # -------------------------

                renderer = DOCX_RENDERERS[
                    renderer_name
                ]

                renderer(
                    doc,
                    value,
                    field_config
                )

    return doc