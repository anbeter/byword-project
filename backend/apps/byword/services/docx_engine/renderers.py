import math
import os

from django.conf import settings

from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def render_text(doc, value, config):
    doc.add_paragraph(str(value))

def render_image(doc, value, config):
    if not value:
        return
    width = config.get("width", 5.5)
    doc.add_picture(
        value.path,
        width=Inches(width)
    )

def render_reference_right(doc, value, config):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(value))
    run.bold = True
    run.font.size = Pt(14)

def render_title_right(doc, value, config):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(value))
    run.bold = True

def render_table_2_columns(doc, value, config):
    if not value:
        return
    rows = math.ceil(len(value) / 2)
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for index, item in enumerate(value):
        col = 0 if index < rows else 1
        row = index if index < rows else index - rows
        table.cell(row, col).text = str(item)

import math


def render_dictionary_table(doc, value, config):
    if not value:
        return
    rows = math.ceil(len(value) / 2)
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for index, item in enumerate(value):
        col = 0 if index < rows else 1
        row = index if index < rows else index - rows
        # ####underline = "_" * ((len(item.translation) * 3) + 2)
        paragraph = table.cell(row,col).paragraphs[0]
        paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(f"{item.verb_en}")
        run.font.size = Pt(14)

import math


def render_scramble_table(doc, value, config):
    if not value:
        return
    if isinstance(value, str):
        words = value.split()
    else:
        words = list(value)
    rows = math.ceil(len(words) / 2)
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for index, word in enumerate(words):
        col = 0 if index < rows else 1
        row = index if index < rows else index - rows
        # ###underline = "_" * ((len(word) * 3) + 2)
        paragraph = table.cell(row,col).paragraphs[0]
        paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT)
        run = paragraph.add_run(f"{word}")
        run.font.size = Pt(14)

def render_wordsearch_image(doc,value,config=None):
    if not value:
        return
    png_path = os.path.join(
        settings.MEDIA_ROOT,
        "png",
        f"{value.id}.png"
    )
    if os.path.exists(png_path):
        doc.add_picture(
            png_path,
            width=Inches(5.8)
        )
