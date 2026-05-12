import random
import re
from docx.shared import Pt


def scramble_sentence_words(sentence):
    punctuation = ""
    if sentence and sentence[-1] in ".?!":
        punctuation = sentence[-1]
        sentence = sentence[:-1]
    words = sentence.split()
    random.shuffle(words)
    result = " ".join(words)
    return f"{result}{punctuation}"


def scramble_sentences(text):
    if not text:
        return ""
    parts = re.split(
        r'([.!?]|\n)',
        text
    )
    sentences = []
    current = ""
    for part in parts:
        if not part:
            continue
        current += part
        if (
            part in ".?!"
            or part == "\n"
        ):
            clean = current.strip()
            if clean:
                sentences.append(
                    scramble_sentence_words(clean)
                )
            current = ""

    if current.strip():
        sentences.append(
            scramble_sentence_words(
                current.strip()
            )
        )

    return "\n".join(sentences)

def render_scramble_sentences(doc, value, config):
    if not value:
        return
    if isinstance(value, str):
        sentences = value.splitlines()
    else:
        sentences = list(value)
    for sentence in sentences:
        clean_sentence = sentence.strip()
        if not clean_sentence:
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(clean_sentence)
        run.font.size = Pt(14)
        # remove espaços para calcular somente letras/sinais
        # size_base = len(clean_sentence.replace(" ", ""))
        size_base = len(clean_sentence)
        underline = "_" * ((size_base * 3) + 2)
        line_paragraph = doc.add_paragraph()
        line_run = line_paragraph.add_run(
            underline
        )
        line_run.font.size = Pt(14)